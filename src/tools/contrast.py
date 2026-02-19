"""WCAG 2.1 AA color contrast checking and fixing.

Pure functions, no file I/O. Uses wcag-contrast-ratio library for
the actual ratio computation per WCAG 2.x algorithm.

Large text threshold: >= 18pt, or >= 14pt bold.
Normal text: 4.5:1 ratio required.
Large text: 3.0:1 ratio required.
"""

from __future__ import annotations

import colorsys
import logging
from dataclasses import dataclass, field
from enum import Enum

import wcag_contrast_ratio as wcag

from docx.document import Document
from docx.shared import RGBColor

from src.models.document import ContrastIssue, ParagraphInfo

logger = logging.getLogger(__name__)

NORMAL_TEXT_RATIO = 4.5
LARGE_TEXT_RATIO = 3.0
LARGE_TEXT_SIZE_PT = 18.0
LARGE_TEXT_BOLD_SIZE_PT = 14.0


class FixStrategy(str, Enum):
    DARKEN_FOREGROUND = "darken_foreground"
    LIGHTEN_BACKGROUND = "lighten_background"


@dataclass
class ContrastResult:
    """Result of a contrast check."""
    ratio: float
    passes: bool
    required_ratio: float
    is_large_text: bool


@dataclass
class ContrastFixResult:
    """Result of a contrast fix attempt."""
    original_color: str
    fixed_color: str
    strategy: FixStrategy
    original_ratio: float
    fixed_ratio: float
    required_ratio: float


def hex_to_rgb(hex_color: str) -> tuple[float, float, float]:
    """Convert hex color string to (r, g, b) floats in 0-1 range.

    Accepts '#RRGGBB' or 'RRGGBB'.
    """
    h = hex_color.lstrip("#")
    if len(h) != 6:
        raise ValueError(f"Invalid hex color: {hex_color!r}")
    r = int(h[0:2], 16) / 255.0
    g = int(h[2:4], 16) / 255.0
    b = int(h[4:6], 16) / 255.0
    return (r, g, b)


def rgb_to_hex(r: float, g: float, b: float) -> str:
    """Convert (r, g, b) floats in 0-1 range to '#RRGGBB' hex string."""
    return "#{:02X}{:02X}{:02X}".format(
        round(r * 255), round(g * 255), round(b * 255)
    )


def is_large_text(font_size_pt: float | None, is_bold: bool = False) -> bool:
    """Determine if text qualifies as 'large text' per WCAG."""
    if font_size_pt is None:
        return False
    if is_bold:
        return font_size_pt >= LARGE_TEXT_BOLD_SIZE_PT
    return font_size_pt >= LARGE_TEXT_SIZE_PT


def check_contrast(
    foreground: str,
    background: str,
    font_size_pt: float | None = None,
    is_bold: bool = False,
) -> ContrastResult:
    """Check if foreground/background colors meet WCAG AA contrast requirements.

    Args:
        foreground: Hex color of text (e.g. '#333333').
        background: Hex color of background (e.g. '#FFFFFF').
        font_size_pt: Font size in points, if known.
        is_bold: Whether text is bold.

    Returns:
        ContrastResult with ratio, pass/fail, and threshold used.
    """
    fg_rgb = hex_to_rgb(foreground)
    bg_rgb = hex_to_rgb(background)
    ratio = wcag.rgb(fg_rgb, bg_rgb)

    large = is_large_text(font_size_pt, is_bold)
    required = LARGE_TEXT_RATIO if large else NORMAL_TEXT_RATIO

    return ContrastResult(
        ratio=round(ratio, 2),
        passes=wcag.passes_AA(ratio, large=large),
        required_ratio=required,
        is_large_text=large,
    )


def fix_contrast(
    foreground: str,
    background: str,
    font_size_pt: float | None = None,
    is_bold: bool = False,
    strategy: FixStrategy = FixStrategy.DARKEN_FOREGROUND,
) -> ContrastFixResult:
    """Fix a contrast issue by adjusting one color via binary search in HLS space.

    Minimizes color change while preserving hue. Only adjusts lightness.

    Args:
        foreground: Hex color of text.
        background: Hex color of background.
        font_size_pt: Font size in points.
        is_bold: Whether text is bold.
        strategy: Which color to adjust.

    Returns:
        ContrastFixResult with original and fixed colors.
    """
    fg_rgb = hex_to_rgb(foreground)
    bg_rgb = hex_to_rgb(background)
    original_ratio = wcag.rgb(fg_rgb, bg_rgb)

    large = is_large_text(font_size_pt, is_bold)
    required = LARGE_TEXT_RATIO if large else NORMAL_TEXT_RATIO

    if strategy == FixStrategy.DARKEN_FOREGROUND:
        color_to_fix = foreground
        fixed_hex = _binary_search_lightness(
            color_to_fix, background, required, darken=True
        )
    else:
        color_to_fix = background
        fixed_hex = _binary_search_lightness(
            color_to_fix, foreground, required, darken=False
        )

    fixed_rgb = hex_to_rgb(fixed_hex)
    if strategy == FixStrategy.DARKEN_FOREGROUND:
        fixed_ratio = wcag.rgb(fixed_rgb, hex_to_rgb(background))
    else:
        fixed_ratio = wcag.rgb(hex_to_rgb(foreground), fixed_rgb)

    return ContrastFixResult(
        original_color=color_to_fix,
        fixed_color=fixed_hex,
        strategy=strategy,
        original_ratio=round(original_ratio, 2),
        fixed_ratio=round(fixed_ratio, 2),
        required_ratio=required,
    )


def _binary_search_lightness(
    color_hex: str,
    other_hex: str,
    target_ratio: float,
    darken: bool,
    max_iterations: int = 30,
) -> str:
    """Binary search in HLS lightness to find a color meeting the target contrast ratio.

    Preserves hue and saturation; only adjusts lightness.
    """
    rgb = hex_to_rgb(color_hex)
    other_rgb = hex_to_rgb(other_hex)

    h, l, s = colorsys.rgb_to_hls(*rgb)

    # Add a small margin so rounding doesn't land us just below the target
    effective_target = target_ratio + 0.05

    if darken:
        lo, hi = 0.0, l  # search darker (lower lightness)
    else:
        lo, hi = l, 1.0  # search lighter (higher lightness)

    best_hex = color_hex

    for _ in range(max_iterations):
        mid = (lo + hi) / 2.0
        candidate_rgb = colorsys.hls_to_rgb(h, mid, s)
        # Clamp to valid range
        candidate_rgb = tuple(max(0.0, min(1.0, c)) for c in candidate_rgb)

        if darken:
            ratio = wcag.rgb(candidate_rgb, other_rgb)
        else:
            ratio = wcag.rgb(other_rgb, candidate_rgb)

        if ratio >= effective_target:
            best_hex = rgb_to_hex(*candidate_rgb)
            if darken:
                lo = mid  # passes — try lighter (closer to original)
            else:
                hi = mid  # passes — try darker (closer to original)
        else:
            if darken:
                hi = mid  # fails — try darker
            else:
                lo = mid  # fails — try lighter

    return best_hex


def analyze_document_contrast(
    paragraphs: list[ParagraphInfo],
    default_bg: str = "#FFFFFF",
) -> list[ContrastIssue]:
    """Analyze all paragraphs for contrast issues.

    Args:
        paragraphs: List of parsed paragraphs with run formatting.
        default_bg: Default background color (usually white).

    Returns:
        List of ContrastIssue for each failing run.
    """
    issues: list[ContrastIssue] = []

    for para in paragraphs:
        for run_idx, run in enumerate(para.runs):
            if not run.text.strip():
                continue

            fg = run.color
            if fg is None:
                fg = "#000000"  # default black text

            # Skip theme colors that we can't resolve
            if not fg.startswith("#"):
                logger.warning(
                    "Skipping unresolvable color %r in %s run %d",
                    fg, para.id, run_idx,
                )
                continue

            bold = run.bold is True
            result = check_contrast(fg, default_bg, run.font_size_pt, bold)

            if not result.passes:
                preview = run.text[:50]
                issues.append(ContrastIssue(
                    paragraph_id=para.id,
                    run_index=run_idx,
                    text_preview=preview,
                    foreground=fg,
                    background=default_bg,
                    contrast_ratio=result.ratio,
                    required_ratio=result.required_ratio,
                    is_large_text=result.is_large_text,
                    font_size_pt=run.font_size_pt,
                    is_bold=bold,
                ))

    return issues


@dataclass
class ApplyContrastResult:
    """Result of applying a contrast fix to a single run."""
    success: bool
    change: str = ""
    error: str = ""


@dataclass
class BulkContrastResult:
    """Result of fixing all contrast issues in a document."""
    success: bool
    fixes_applied: int = 0
    fixes_failed: int = 0
    changes: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


def apply_contrast_fix(
    doc: Document,
    paragraph_index: int,
    run_index: int,
    new_color: str,
) -> ApplyContrastResult:
    """Apply a contrast fix to a single run in the document.

    Args:
        doc: python-docx Document (opened for modification).
        paragraph_index: Index of the paragraph containing the run.
        run_index: Index of the run within the paragraph.
        new_color: Hex color to set (e.g. '#4A4A4A').

    Returns:
        ApplyContrastResult with success/failure.
    """
    try:
        if paragraph_index >= len(doc.paragraphs):
            return ApplyContrastResult(
                success=False,
                error=f"Paragraph index {paragraph_index} out of range (doc has {len(doc.paragraphs)} paragraphs)",
            )

        para = doc.paragraphs[paragraph_index]
        if run_index >= len(para.runs):
            return ApplyContrastResult(
                success=False,
                error=f"Run index {run_index} out of range (paragraph has {len(para.runs)} runs)",
            )

        run = para.runs[run_index]
        old_color = f"#{run.font.color.rgb}" if run.font.color and run.font.color.rgb else "default"

        # Parse hex and set RGBColor
        h = new_color.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        preview = run.text[:40]
        change = f"p{paragraph_index} run{run_index}: {old_color} -> {new_color} ({preview!r})"
        logger.info("Contrast fix: %s", change)
        return ApplyContrastResult(success=True, change=change)

    except Exception as e:
        return ApplyContrastResult(success=False, error=f"Failed to apply contrast fix: {e}")


def fix_all_document_contrast(
    doc: Document,
    paragraphs: list[ParagraphInfo],
    default_bg: str = "#FFFFFF",
) -> BulkContrastResult:
    """Analyze and fix all contrast issues in a document in one pass.

    Finds every run that fails WCAG AA contrast, computes a minimal fix
    (darken foreground), and applies it to the python-docx document.

    Args:
        doc: python-docx Document (opened for modification).
        paragraphs: Parsed paragraph models (for analysis).
        default_bg: Default background color.

    Returns:
        BulkContrastResult with counts and change log.
    """
    issues = analyze_document_contrast(paragraphs, default_bg)

    if not issues:
        return BulkContrastResult(success=True, changes=["No contrast issues found"])

    fixes_applied = 0
    fixes_failed = 0
    changes: list[str] = []
    errors: list[str] = []

    for issue in issues:
        # Compute the fixed color
        fix_result = fix_contrast(
            issue.foreground,
            issue.background,
            issue.font_size_pt,
            issue.is_bold,
            strategy=FixStrategy.DARKEN_FOREGROUND,
        )

        # Find the paragraph index from paragraph_id (p_0 -> 0)
        para_idx = int(issue.paragraph_id.split("_")[1])

        # Apply the fix
        apply_result = apply_contrast_fix(doc, para_idx, issue.run_index, fix_result.fixed_color)

        if apply_result.success:
            fixes_applied += 1
            changes.append(apply_result.change)
        else:
            fixes_failed += 1
            errors.append(apply_result.error)

    logger.info(
        "Bulk contrast fix: %d applied, %d failed out of %d issues",
        fixes_applied, fixes_failed, len(issues),
    )

    return BulkContrastResult(
        success=fixes_failed == 0,
        fixes_applied=fixes_applied,
        fixes_failed=fixes_failed,
        changes=changes,
        errors=errors,
    )
