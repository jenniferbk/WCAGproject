"""Parse .docx files into a DocumentModel.

Uses python-docx for high-level access and lxml for raw XML when needed.
Two-pass approach for fake heading detection:
  1. First pass: collect all font sizes and paragraph data
  2. Second pass: score fake heading candidates against the median font size
"""

from __future__ import annotations

import logging
import re
import statistics
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.models.document import (
    CellInfo,
    ContentOrderItem,
    ContentType,
    DocumentModel,
    DocumentStats,
    FakeHeadingSignals,
    ImageInfo,
    LinkInfo,
    MetadataInfo,
    ParagraphInfo,
    RunInfo,
    TableInfo,
)
from src.tools.image_extract import extract_images_from_docx

logger = logging.getLogger(__name__)

# Default font size in points when nothing else is specified
DEFAULT_FONT_SIZE_PT = 12.0

# Fake heading scoring weights
_WEIGHT_ALL_BOLD = 0.30
_WEIGHT_FONT_SIZE = 0.25
_WEIGHT_SHORT = 0.20
_WEIGHT_FOLLOWED_BY = 0.15
_WEIGHT_NOT_IN_TABLE = 0.10


@dataclass
class ParseResult:
    """Result of parsing a docx file."""
    success: bool
    document: DocumentModel | None = None
    error: str = ""
    warnings: list[str] = field(default_factory=list)


def parse_docx(filepath: str | Path) -> ParseResult:
    """Parse a .docx file into a DocumentModel.

    Args:
        filepath: Path to the .docx file.

    Returns:
        ParseResult with success/failure and the DocumentModel.
    """
    filepath = Path(filepath)
    warnings: list[str] = []

    if not filepath.exists():
        return ParseResult(success=False, error=f"File not found: {filepath}")
    if not filepath.suffix.lower() == ".docx":
        return ParseResult(success=False, error=f"Not a .docx file: {filepath}")

    try:
        doc = docx.Document(str(filepath))
    except Exception as e:
        return ParseResult(success=False, error=f"Failed to open docx: {e}")

    try:
        # Extract metadata
        metadata = _extract_metadata(doc)

        # Extract images first (we need their IDs when parsing paragraphs)
        paragraphs_text = [p.text for p in doc.paragraphs]
        image_result = extract_images_from_docx(doc, paragraphs_text)
        images = image_result.images
        warnings.extend(image_result.warnings)

        # Build image lookup: paragraph_id -> list of image IDs
        para_image_map: dict[str, list[str]] = {}
        for img in images:
            para_image_map.setdefault(img.paragraph_id, []).append(img.id)

        # Resolve default font size from document styles
        default_font_size = _get_default_font_size(doc)

        # First pass: parse paragraphs and tables, collect font sizes
        all_font_sizes: list[float] = []
        paragraphs: list[ParagraphInfo] = []
        tables: list[TableInfo] = []
        links: list[LinkInfo] = []
        content_order: list[ContentOrderItem] = []

        para_counter = 0
        table_counter = 0
        link_counter = 0

        for block in doc.iter_inner_content():
            if isinstance(block, Paragraph):
                para_id = f"p_{para_counter}"
                para_info, para_links, para_fonts = _parse_paragraph(
                    block, para_id, link_counter, default_font_size,
                    para_image_map.get(para_id, []),
                )
                paragraphs.append(para_info)
                links.extend(para_links)
                link_counter += len(para_links)
                all_font_sizes.extend(para_fonts)
                content_order.append(ContentOrderItem(
                    content_type=ContentType.PARAGRAPH, id=para_id
                ))
                para_counter += 1

            elif isinstance(block, Table):
                tbl_id = f"tbl_{table_counter}"
                tbl_info = _parse_table(block, tbl_id)
                tables.append(tbl_info)
                content_order.append(ContentOrderItem(
                    content_type=ContentType.TABLE, id=tbl_id
                ))
                table_counter += 1

        # Second pass: compute fake heading signals
        median_font = (
            statistics.median(all_font_sizes) if all_font_sizes
            else default_font_size
        )
        paragraphs = _score_fake_headings(paragraphs, median_font)

        # Compute stats
        heading_count = sum(
            1 for p in paragraphs if p.heading_level is not None
        )
        images_missing_alt = sum(1 for img in images if not img.alt_text)
        fake_heading_candidates = sum(
            1 for p in paragraphs
            if p.fake_heading_signals is not None and p.fake_heading_signals.score >= 0.5
        )

        stats = DocumentStats(
            paragraph_count=len(paragraphs),
            table_count=len(tables),
            image_count=len(images),
            link_count=len(links),
            heading_count=heading_count,
            images_missing_alt=images_missing_alt,
            fake_heading_candidates=fake_heading_candidates,
        )

        document = DocumentModel(
            source_format="docx",
            source_path=str(filepath),
            metadata=metadata,
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            links=links,
            content_order=content_order,
            stats=stats,
            parse_warnings=warnings,
        )

        return ParseResult(success=True, document=document, warnings=warnings)

    except Exception as e:
        logger.exception("Error parsing docx: %s", filepath)
        return ParseResult(success=False, error=f"Parse error: {e}", warnings=warnings)


def _extract_metadata(doc: Document) -> MetadataInfo:
    """Extract document metadata from core properties."""
    props = doc.core_properties
    return MetadataInfo(
        title=props.title or "",
        author=props.author or "",
        language=props.language or "",
        subject=props.subject or "",
        created=str(props.created) if props.created else "",
        modified=str(props.modified) if props.modified else "",
    )


def _get_default_font_size(doc: Document) -> float:
    """Get the default font size from document styles, falling back to 12pt."""
    try:
        default_style = doc.styles.default(1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        if default_style and default_style.font and default_style.font.size:
            return default_style.font.size.pt
    except Exception:
        pass
    return DEFAULT_FONT_SIZE_PT


def _resolve_font_size(
    run: docx.text.run.Run,
    paragraph: Paragraph,
    default_font_size: float,
) -> float | None:
    """Resolve font size through the inheritance chain: run -> paragraph style -> default.

    Returns None only if there is truly no font size information.
    """
    # Direct run formatting
    if run.font.size is not None:
        return run.font.size.pt

    # Paragraph style
    try:
        style = paragraph.style
        if style and style.font and style.font.size:
            return style.font.size.pt
    except Exception:
        pass

    return default_font_size


def _parse_paragraph(
    paragraph: Paragraph,
    para_id: str,
    link_counter_start: int,
    default_font_size: float,
    image_ids: list[str],
) -> tuple[ParagraphInfo, list[LinkInfo], list[float]]:
    """Parse a single paragraph into ParagraphInfo.

    Returns:
        Tuple of (ParagraphInfo, list of links, list of font sizes found).
    """
    runs: list[RunInfo] = []
    links: list[LinkInfo] = []
    font_sizes: list[float] = []
    link_counter = link_counter_start

    # Parse runs
    for run in paragraph.runs:
        font_size = _resolve_font_size(run, paragraph, default_font_size)
        if font_size is not None:
            font_sizes.append(font_size)

        # Resolve color
        color_hex = None
        if run.font.color and run.font.color.rgb:
            color_hex = f"#{run.font.color.rgb}"

        runs.append(RunInfo(
            text=run.text,
            bold=run.bold,
            italic=run.italic,
            underline=run.underline,
            font_size_pt=font_size,
            font_name=run.font.name,
            color=color_hex,
        ))

    # Parse hyperlinks from XML
    para_elem = paragraph._element
    for hyperlink in para_elem.findall(qn("w:hyperlink")):
        r_id = hyperlink.get(qn("r:id"))
        url = ""
        if r_id:
            try:
                rel = paragraph.part.rels.get(r_id)
                if rel:
                    url = rel._target
            except Exception:
                pass

        # Get hyperlink text from its w:r/w:t elements
        link_text_parts = []
        for r_elem in hyperlink.findall(qn("w:r")):
            for t_elem in r_elem.findall(qn("w:t")):
                if t_elem.text:
                    link_text_parts.append(t_elem.text)
        link_text = "".join(link_text_parts)

        if link_text or url:
            link_id = f"link_{link_counter}"
            links.append(LinkInfo(
                id=link_id,
                text=link_text,
                url=url,
                paragraph_id=para_id,
            ))
            link_counter += 1

    # Determine heading level
    heading_level = None
    style_name = paragraph.style.name if paragraph.style else "Normal"
    heading_match = re.match(r"Heading\s*(\d+)", style_name)
    if heading_match:
        heading_level = int(heading_match.group(1))

    # Detect list items
    is_list_item = False
    list_level = None
    num_pr = para_elem.find(f".//{qn('w:numPr')}")
    if num_pr is not None:
        is_list_item = True
        ilvl = num_pr.find(qn("w:ilvl"))
        if ilvl is not None:
            list_level = int(ilvl.get(qn("w:val"), "0"))

    # Alignment
    alignment = None
    if paragraph.alignment is not None:
        alignment_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        alignment = alignment_map.get(paragraph.alignment, str(paragraph.alignment))

    para_info = ParagraphInfo(
        id=para_id,
        text=paragraph.text,
        style_name=style_name,
        heading_level=heading_level,
        runs=runs,
        links=links,
        image_ids=image_ids,
        alignment=alignment,
        is_list_item=is_list_item,
        list_level=list_level,
    )

    return para_info, links, font_sizes


def _parse_table(table: Table, tbl_id: str) -> TableInfo:
    """Parse a table into TableInfo."""
    rows: list[list[CellInfo]] = []
    header_row_count = 0

    for row_idx, row in enumerate(table.rows):
        cells: list[CellInfo] = []

        # Check if row is marked as a header row
        tr_pr = row._tr.find(qn("w:trPr"))
        if tr_pr is not None:
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is not None:
                header_row_count = max(header_row_count, row_idx + 1)

        for cell in row.cells:
            cell_paras = [p.text for p in cell.paragraphs]
            cell_text = "\n".join(cell_paras)

            # Grid span (horizontal merge)
            grid_span = 1
            tc_pr = cell._tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                gs = tc_pr.find(qn("w:gridSpan"))
                if gs is not None:
                    grid_span = int(gs.get(qn("w:val"), "1"))

                # Vertical merge
                v_merge_elem = tc_pr.find(qn("w:vMerge"))
                v_merge = None
                if v_merge_elem is not None:
                    v_merge = v_merge_elem.get(qn("w:val"), "continue")
            else:
                v_merge = None

            cells.append(CellInfo(
                text=cell_text,
                paragraphs=cell_paras,
                grid_span=grid_span,
                v_merge=v_merge,
            ))

        rows.append(cells)

    style_name = table.style.name if table.style else ""
    row_count = len(rows)
    col_count = max((len(r) for r in rows), default=0)

    return TableInfo(
        id=tbl_id,
        rows=rows,
        header_row_count=header_row_count,
        has_header_style="Header" in style_name or "header" in style_name,
        style_name=style_name,
        row_count=row_count,
        col_count=col_count,
    )


def _score_fake_headings(
    paragraphs: list[ParagraphInfo],
    median_font_size: float,
) -> list[ParagraphInfo]:
    """Score paragraphs for fake heading likelihood.

    Only scores paragraphs with Normal style that have text.
    Uses weighted combination of heuristic signals.
    """
    result: list[ParagraphInfo] = []

    for i, para in enumerate(paragraphs):
        # Skip headings, empty paragraphs, and list items
        if para.heading_level is not None or not para.text.strip() or para.is_list_item:
            result.append(para)
            continue

        # Only check Normal-style paragraphs (or similar body styles)
        if para.style_name not in ("Normal", "Body Text", "Body", "Default Paragraph Font"):
            result.append(para)
            continue

        # Check if all runs are bold
        text_runs = [r for r in para.runs if r.text.strip()]
        if not text_runs:
            result.append(para)
            continue

        all_bold = all(r.bold is True for r in text_runs)

        # Get font size (use max across runs)
        run_sizes = [r.font_size_pt for r in text_runs if r.font_size_pt is not None]
        max_font_size = max(run_sizes) if run_sizes else None
        font_above_avg = (
            max_font_size is not None and max_font_size >= median_font_size + 2.0
        )

        # Short text (< ~10 words)
        word_count = len(para.text.split())
        is_short = word_count < 10

        # Followed by non-bold text
        followed_by_non_bold = False
        for j in range(i + 1, min(i + 3, len(paragraphs))):
            next_para = paragraphs[j]
            if next_para.text.strip():
                next_text_runs = [r for r in next_para.runs if r.text.strip()]
                if next_text_runs:
                    followed_by_non_bold = not all(
                        r.bold is True for r in next_text_runs
                    )
                break

        # Must have at least bold to be a candidate
        if not all_bold:
            result.append(para)
            continue

        # Compute weighted score
        score = (
            _WEIGHT_ALL_BOLD * (1.0 if all_bold else 0.0)
            + _WEIGHT_FONT_SIZE * (1.0 if font_above_avg else 0.0)
            + _WEIGHT_SHORT * (1.0 if is_short else 0.0)
            + _WEIGHT_FOLLOWED_BY * (1.0 if followed_by_non_bold else 0.0)
            + _WEIGHT_NOT_IN_TABLE * 1.0  # always True for paragraphs
        )

        signals = FakeHeadingSignals(
            all_runs_bold=all_bold,
            font_size_pt=max_font_size,
            font_size_above_avg=font_above_avg,
            is_short=is_short,
            followed_by_non_bold=followed_by_non_bold,
            not_in_table=True,
            score=round(score, 3),
        )

        # Reconstruct with signals (frozen model)
        updated = para.model_copy(update={"fake_heading_signals": signals})
        result.append(updated)

    return result
