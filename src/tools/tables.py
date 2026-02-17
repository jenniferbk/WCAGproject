"""Table header marking and structure analysis for .docx files.

Uses raw XML manipulation to set w:tblHeader on header rows since
python-docx doesn't expose this property.

Targets:
- 1.3.1: Table structure uses semantic markup (header rows marked)
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.models.document import TableInfo

logger = logging.getLogger(__name__)


@dataclass
class TableResult:
    """Result of a table operation."""
    success: bool
    changes: list[str] = field(default_factory=list)
    error: str = ""


@dataclass
class TableAnalysis:
    """Analysis of table structure for accessibility."""
    table_id: str
    row_count: int
    col_count: int
    has_header_row: bool
    header_row_count: int
    has_merged_cells: bool
    is_complex: bool  # merged cells or nested tables
    issues: list[str] = field(default_factory=list)


def analyze_table(table_info: TableInfo) -> TableAnalysis:
    """Analyze a parsed table for accessibility issues.

    Args:
        table_info: TableInfo from the DocumentModel.

    Returns:
        TableAnalysis with structure details and issues.
    """
    issues: list[str] = []

    # Check for header row
    if table_info.header_row_count == 0:
        issues.append("Table has no marked header row")

    # Check for merged cells
    has_merged = False
    for row in table_info.rows:
        for cell in row:
            if cell.grid_span > 1 or cell.v_merge is not None:
                has_merged = True
                break
        if has_merged:
            break

    if has_merged:
        issues.append("Table has merged cells — may need manual review")

    is_complex = has_merged

    return TableAnalysis(
        table_id=table_info.id,
        row_count=table_info.row_count,
        col_count=table_info.col_count,
        has_header_row=table_info.header_row_count > 0,
        header_row_count=table_info.header_row_count,
        has_merged_cells=has_merged,
        is_complex=is_complex,
        issues=issues,
    )


def mark_header_rows(
    doc: Document,
    table_index: int,
    header_count: int = 1,
) -> TableResult:
    """Mark the first N rows of a table as header rows via raw XML.

    Sets w:tblHeader on each header row's w:trPr element so that
    headers repeat across pages and are recognized by assistive technology.

    Args:
        doc: python-docx Document.
        table_index: Index of the table in doc.tables.
        header_count: Number of rows to mark as headers (default 1).

    Returns:
        TableResult with success/failure.
    """
    try:
        if table_index >= len(doc.tables):
            return TableResult(
                success=False,
                error=f"Table index {table_index} out of range (doc has {len(doc.tables)} tables)",
            )

        table = doc.tables[table_index]

        if header_count > len(table.rows):
            return TableResult(
                success=False,
                error=f"header_count {header_count} exceeds table row count {len(table.rows)}",
            )

        changes: list[str] = []

        for row_idx in range(header_count):
            row = table.rows[row_idx]
            tr_pr = row._tr.get_or_add_trPr()

            # Check if already marked
            existing = tr_pr.find(qn("w:tblHeader"))
            if existing is not None:
                changes.append(f"tbl_{table_index} row {row_idx}: already marked as header")
                continue

            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
            changes.append(f"tbl_{table_index} row {row_idx}: marked as header row")
            logger.info("Marked table %d row %d as header", table_index, row_idx)

        return TableResult(success=True, changes=changes)

    except Exception as e:
        return TableResult(success=False, error=f"Failed to mark header rows: {e}")


def remove_header_rows(
    doc: Document,
    table_index: int,
) -> TableResult:
    """Remove header row marking from all rows in a table.

    Args:
        doc: python-docx Document.
        table_index: Index of the table.

    Returns:
        TableResult with success/failure.
    """
    try:
        if table_index >= len(doc.tables):
            return TableResult(
                success=False,
                error=f"Table index {table_index} out of range",
            )

        table = doc.tables[table_index]
        changes: list[str] = []

        for row_idx, row in enumerate(table.rows):
            tr_pr = row._tr.find(qn("w:trPr"))
            if tr_pr is not None:
                tbl_header = tr_pr.find(qn("w:tblHeader"))
                if tbl_header is not None:
                    tr_pr.remove(tbl_header)
                    changes.append(f"tbl_{table_index} row {row_idx}: header marking removed")

        if not changes:
            changes.append(f"tbl_{table_index}: no header rows to remove")

        return TableResult(success=True, changes=changes)

    except Exception as e:
        return TableResult(success=False, error=f"Failed to remove header rows: {e}")


def analyze_all_tables(tables: list[TableInfo]) -> list[TableAnalysis]:
    """Analyze all tables in the document for accessibility.

    Args:
        tables: List of TableInfo from the DocumentModel.

    Returns:
        List of TableAnalysis for each table.
    """
    return [analyze_table(t) for t in tables]
