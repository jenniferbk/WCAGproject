"""Shared test fixtures: programmatically build .docx files for testing."""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image


def _make_test_image(width: int = 100, height: int = 50, color: str = "red") -> bytes:
    """Create a simple test image and return as PNG bytes."""
    img = Image.new("RGB", (width, height), color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    """A simple docx with title, headings, body text, and a list."""
    doc = Document()
    doc.core_properties.title = "Test Document"
    doc.core_properties.author = "Test Author"
    doc.core_properties.language = "en-US"

    doc.add_heading("Main Title", level=1)
    doc.add_paragraph("This is the first body paragraph with some text.")
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("Body text under section one.")
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Deeper content here.")

    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def fake_headings_docx(tmp_path: Path) -> Path:
    """A docx with fake headings: bold, larger text that SHOULD be headings but aren't."""
    doc = Document()

    # Real heading
    doc.add_heading("Real Heading", level=1)
    doc.add_paragraph("Normal text after real heading.")

    # Fake heading: bold, larger font, short, Normal style
    fake = doc.add_paragraph()
    run = fake.add_run("Fake Section Title")
    run.bold = True
    run.font.size = Pt(16)
    # Follow with normal text
    doc.add_paragraph("This is normal body text that follows the fake heading.")

    # Another fake heading
    fake2 = doc.add_paragraph()
    run2 = fake2.add_run("Another Bold Title")
    run2.bold = True
    run2.font.size = Pt(14)
    doc.add_paragraph("More body text here.")

    # NOT a fake heading: bold but long text
    not_fake = doc.add_paragraph()
    run3 = not_fake.add_run(
        "This is a long bold paragraph that has way too many words to be a heading "
        "because headings are typically short and descriptive."
    )
    run3.bold = True
    run3.font.size = Pt(12)

    path = tmp_path / "fake_headings.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def image_docx(tmp_path: Path) -> Path:
    """A docx with an inline image that has alt text."""
    doc = Document()
    doc.add_paragraph("Text before the image.")

    # Add an image
    img_bytes = _make_test_image(200, 100, "blue")
    img_path = tmp_path / "test_image.png"
    img_path.write_bytes(img_bytes)
    doc.add_picture(str(img_path), width=Pt(200))

    # Set alt text via XML
    # The picture was added to the last paragraph
    last_para = doc.paragraphs[-1]
    for drawing in last_para._element.findall(
        f".//{{{qn('w:drawing').split('}')[0][1:]}}}"
        if False else f".//{qn('w:drawing')}"
    ):
        inline = drawing.find(
            f"{{{qn('wp:inline').split('}')[0][1:]}}}"
            if False else qn("wp:inline")
        )
        if inline is not None:
            doc_pr = inline.find(qn("wp:docPr"))
            if doc_pr is not None:
                doc_pr.set("descr", "A blue test rectangle")

    doc.add_paragraph("Text after the image.")

    path = tmp_path / "with_image.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def image_no_alt_docx(tmp_path: Path) -> Path:
    """A docx with an image that has NO alt text."""
    doc = Document()
    doc.add_paragraph("Before image.")

    img_bytes = _make_test_image(150, 75, "green")
    img_path = tmp_path / "no_alt_image.png"
    img_path.write_bytes(img_bytes)
    doc.add_picture(str(img_path), width=Pt(150))

    doc.add_paragraph("After image.")

    path = tmp_path / "no_alt_image.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def table_docx(tmp_path: Path) -> Path:
    """A docx with a table including header row."""
    doc = Document()
    doc.add_paragraph("Table below:")

    table = doc.add_table(rows=3, cols=3)
    table.style = doc.styles["Table Grid"]

    # Header row
    for i, text in enumerate(["Name", "Age", "City"]):
        table.rows[0].cells[i].text = text

    # Mark first row as header
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)

    # Data rows
    for i, (name, age, city) in enumerate([("Alice", "30", "NYC"), ("Bob", "25", "LA")]):
        table.rows[i + 1].cells[0].text = name
        table.rows[i + 1].cells[1].text = age
        table.rows[i + 1].cells[2].text = city

    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def contrast_docx(tmp_path: Path) -> Path:
    """A docx with colored text for contrast testing."""
    doc = Document()

    # Good contrast: black on white (default)
    doc.add_paragraph("Black text on white background.")

    # Poor contrast: light gray text
    para = doc.add_paragraph()
    run = para.add_run("Light gray text that fails contrast.")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.size = Pt(12)

    # Large text with mediocre contrast
    para2 = doc.add_paragraph()
    run2 = para2.add_run("Large gray text.")
    run2.font.color.rgb = RGBColor(0x76, 0x76, 0x76)
    run2.font.size = Pt(20)

    path = tmp_path / "contrast.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def no_metadata_docx(tmp_path: Path) -> Path:
    """A docx with no title or language set."""
    doc = Document()
    doc.add_paragraph("A document with no metadata.")
    path = tmp_path / "no_metadata.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def fake_lists_docx(tmp_path: Path) -> Path:
    """A docx with fake lists: manual numbering and bullet characters."""
    doc = Document()
    doc.add_paragraph("Here is a list of items:")

    # Fake numbered list
    doc.add_paragraph("1. First item in the list")
    doc.add_paragraph("2. Second item in the list")
    doc.add_paragraph("3. Third item in the list")

    doc.add_paragraph("And some bullet points:")

    # Fake bulleted list
    doc.add_paragraph("- Bullet one")
    doc.add_paragraph("- Bullet two")
    doc.add_paragraph("- Bullet three")

    doc.add_paragraph("End of lists.")

    path = tmp_path / "fake_lists.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def table_no_header_docx(tmp_path: Path) -> Path:
    """A docx with a table that has NO marked header row."""
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.style = doc.styles["Table Grid"]
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Alpha"
    table.rows[1].cells[1].text = "100"
    table.rows[2].cells[0].text = "Beta"
    table.rows[2].cells[1].text = "200"

    path = tmp_path / "table_no_header.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def skipped_headings_docx(tmp_path: Path) -> Path:
    """A docx with skipped heading levels (H1 -> H3, no H2)."""
    doc = Document()
    doc.core_properties.title = "Skipped Headings"
    doc.core_properties.language = "en"

    doc.add_heading("Title", level=1)
    doc.add_paragraph("Intro text.")
    doc.add_heading("Deep Section", level=3)  # skips H2
    doc.add_paragraph("Content.")

    path = tmp_path / "skipped_headings.docx"
    doc.save(str(path))
    return path
