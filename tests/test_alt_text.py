"""Tests for alt text reading and writing."""

from pathlib import Path

import docx
import pytest

from src.tools.alt_text import get_all_alt_text, set_alt_text, set_decorative


class TestGetAltText:
    def test_image_with_alt_text(self, image_docx: Path):
        doc = docx.Document(str(image_docx))
        images = get_all_alt_text(doc)
        assert len(images) == 1
        assert images[0].has_alt_text
        assert images[0].current_alt_text == "A blue test rectangle"

    def test_image_without_alt_text(self, image_no_alt_docx: Path):
        doc = docx.Document(str(image_no_alt_docx))
        images = get_all_alt_text(doc)
        assert len(images) == 1
        assert not images[0].has_alt_text
        assert images[0].current_alt_text == ""


class TestSetAltText:
    def test_set_alt_text(self, image_no_alt_docx: Path):
        doc = docx.Document(str(image_no_alt_docx))

        # Find which paragraph has the image
        images = get_all_alt_text(doc)
        para_idx = images[0].paragraph_index

        result = set_alt_text(doc, para_idx, "A green test image")
        assert result.success
        assert len(result.changes) == 1

        # Verify it was set
        images_after = get_all_alt_text(doc)
        assert images_after[0].current_alt_text == "A green test image"

    def test_overwrite_alt_text(self, image_docx: Path):
        doc = docx.Document(str(image_docx))
        images = get_all_alt_text(doc)
        para_idx = images[0].paragraph_index

        result = set_alt_text(doc, para_idx, "Updated description")
        assert result.success

        images_after = get_all_alt_text(doc)
        assert images_after[0].current_alt_text == "Updated description"

    def test_invalid_paragraph_index(self, image_docx: Path):
        doc = docx.Document(str(image_docx))
        result = set_alt_text(doc, 999, "text")
        assert not result.success
        assert "out of range" in result.error

    def test_invalid_drawing_index(self, image_docx: Path):
        doc = docx.Document(str(image_docx))
        images = get_all_alt_text(doc)
        para_idx = images[0].paragraph_index
        result = set_alt_text(doc, para_idx, "text", drawing_index=99)
        assert not result.success
        assert "out of range" in result.error


class TestSetDecorative:
    def test_mark_decorative(self, image_docx: Path):
        doc = docx.Document(str(image_docx))
        images = get_all_alt_text(doc)
        para_idx = images[0].paragraph_index

        result = set_decorative(doc, para_idx)
        assert result.success

        images_after = get_all_alt_text(doc)
        assert images_after[0].current_alt_text == ""


class TestPersistence:
    def test_alt_text_survives_save(self, image_no_alt_docx: Path, tmp_path: Path):
        doc = docx.Document(str(image_no_alt_docx))
        images = get_all_alt_text(doc)
        para_idx = images[0].paragraph_index

        set_alt_text(doc, para_idx, "Persistent alt text")

        output_path = tmp_path / "saved.docx"
        doc.save(str(output_path))

        # Reopen and verify
        doc2 = docx.Document(str(output_path))
        images2 = get_all_alt_text(doc2)
        assert images2[0].current_alt_text == "Persistent alt text"
