"""Tests for contrast checking and fixing."""

import pytest

from src.models.document import ParagraphInfo, RunInfo
from src.tools.contrast import (
    FixStrategy,
    analyze_document_contrast,
    apply_contrast_fix,
    check_contrast,
    fix_all_document_contrast,
    fix_contrast,
    hex_to_rgb,
    is_large_text,
    rgb_to_hex,
)


class TestHexConversion:
    def test_hex_to_rgb_black(self):
        assert hex_to_rgb("#000000") == (0.0, 0.0, 0.0)

    def test_hex_to_rgb_white(self):
        assert hex_to_rgb("#FFFFFF") == (1.0, 1.0, 1.0)

    def test_hex_to_rgb_red(self):
        r, g, b = hex_to_rgb("#FF0000")
        assert r == 1.0
        assert g == 0.0
        assert b == 0.0

    def test_hex_to_rgb_no_hash(self):
        assert hex_to_rgb("FF0000") == hex_to_rgb("#FF0000")

    def test_rgb_to_hex(self):
        assert rgb_to_hex(1.0, 0.0, 0.0) == "#FF0000"

    def test_rgb_to_hex_black(self):
        assert rgb_to_hex(0.0, 0.0, 0.0) == "#000000"

    def test_roundtrip(self):
        original = "#3A7BD5"
        r, g, b = hex_to_rgb(original)
        result = rgb_to_hex(r, g, b)
        assert result == original

    def test_invalid_hex(self):
        with pytest.raises(ValueError):
            hex_to_rgb("#ZZZ")


class TestLargeText:
    def test_normal_text_not_large(self):
        assert not is_large_text(12.0, is_bold=False)

    def test_18pt_is_large(self):
        assert is_large_text(18.0, is_bold=False)

    def test_14pt_bold_is_large(self):
        assert is_large_text(14.0, is_bold=True)

    def test_14pt_not_bold_not_large(self):
        assert not is_large_text(14.0, is_bold=False)

    def test_none_size_not_large(self):
        assert not is_large_text(None, is_bold=True)


class TestCheckContrast:
    def test_black_on_white_passes(self):
        result = check_contrast("#000000", "#FFFFFF")
        assert result.passes
        assert result.ratio == 21.0

    def test_white_on_white_fails(self):
        result = check_contrast("#FFFFFF", "#FFFFFF")
        assert not result.passes
        assert result.ratio == 1.0

    def test_light_gray_on_white_fails(self):
        result = check_contrast("#AAAAAA", "#FFFFFF")
        assert not result.passes
        assert result.required_ratio == 4.5

    def test_medium_gray_on_white(self):
        # #767676 is roughly the boundary for 4.5:1 on white
        result = check_contrast("#767676", "#FFFFFF")
        assert result.ratio >= 4.5
        assert result.passes

    def test_large_text_lower_threshold(self):
        # A color that fails normal (4.5:1) but passes large (3.0:1)
        result = check_contrast("#999999", "#FFFFFF", font_size_pt=20.0)
        assert result.is_large_text
        assert result.required_ratio == 3.0

    def test_bold_14pt_is_large(self):
        result = check_contrast("#999999", "#FFFFFF", font_size_pt=14.0, is_bold=True)
        assert result.is_large_text


class TestFixContrast:
    def test_darken_foreground(self):
        result = fix_contrast("#AAAAAA", "#FFFFFF", strategy=FixStrategy.DARKEN_FOREGROUND)
        assert result.fixed_ratio >= 4.5
        # Fixed color should be darker (lower hex values)
        assert result.fixed_color != result.original_color

    def test_lighten_background(self):
        # Dark text on medium background
        result = fix_contrast("#333333", "#555555", strategy=FixStrategy.LIGHTEN_BACKGROUND)
        assert result.fixed_ratio >= 4.5

    def test_hue_preservation(self):
        """Fix should preserve the hue of the original color."""
        import colorsys

        original = "#3A7BD5"
        result = fix_contrast(original, "#FFFFFF", strategy=FixStrategy.DARKEN_FOREGROUND)

        orig_r, orig_g, orig_b = hex_to_rgb(original)
        fixed_r, fixed_g, fixed_b = hex_to_rgb(result.fixed_color)

        orig_h, _, _ = colorsys.rgb_to_hls(orig_r, orig_g, orig_b)
        fixed_h, _, _ = colorsys.rgb_to_hls(fixed_r, fixed_g, fixed_b)

        # Hue should be preserved (within floating point tolerance)
        assert abs(orig_h - fixed_h) < 0.01

    def test_already_passing(self):
        result = fix_contrast("#000000", "#FFFFFF", strategy=FixStrategy.DARKEN_FOREGROUND)
        # Should still return a valid result
        assert result.fixed_ratio >= 4.5


class TestAnalyzeDocumentContrast:
    def test_finds_issues(self):
        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="Low contrast text",
                runs=[
                    RunInfo(text="Low contrast text", color="#AAAAAA", font_size_pt=12.0),
                ],
            ),
        ]
        issues = analyze_document_contrast(paragraphs)
        assert len(issues) == 1
        assert issues[0].paragraph_id == "p_0"
        assert issues[0].foreground == "#AAAAAA"

    def test_no_issues_for_good_contrast(self):
        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="Good contrast",
                runs=[
                    RunInfo(text="Good contrast", color="#000000", font_size_pt=12.0),
                ],
            ),
        ]
        issues = analyze_document_contrast(paragraphs)
        assert len(issues) == 0

    def test_default_color_black(self):
        """Runs with no color specified should default to black (passes on white)."""
        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="Default color",
                runs=[
                    RunInfo(text="Default color", font_size_pt=12.0),
                ],
            ),
        ]
        issues = analyze_document_contrast(paragraphs)
        assert len(issues) == 0

    def test_empty_runs_skipped(self):
        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="",
                runs=[
                    RunInfo(text="   ", color="#AAAAAA"),
                ],
            ),
        ]
        issues = analyze_document_contrast(paragraphs)
        assert len(issues) == 0


class TestApplyContrastFix:
    def _make_doc_with_gray_text(self):
        """Create a docx Document with a gray text run for testing."""
        import docx
        from docx.shared import RGBColor as RGB

        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run("Gray text")
        run.font.color.rgb = RGB(0xAA, 0xAA, 0xAA)
        return doc

    def test_apply_fixes_color(self):
        doc = self._make_doc_with_gray_text()
        result = apply_contrast_fix(doc, 0, 0, "#4A4A4A")
        assert result.success
        assert "4A4A4A" in result.change
        # Verify color was actually changed
        actual = f"#{doc.paragraphs[0].runs[0].font.color.rgb}"
        assert actual == "#4A4A4A"

    def test_apply_out_of_range_paragraph(self):
        doc = self._make_doc_with_gray_text()
        result = apply_contrast_fix(doc, 99, 0, "#000000")
        assert not result.success
        assert "out of range" in result.error

    def test_apply_out_of_range_run(self):
        doc = self._make_doc_with_gray_text()
        result = apply_contrast_fix(doc, 0, 99, "#000000")
        assert not result.success
        assert "out of range" in result.error


class TestFixAllDocumentContrast:
    def _make_doc_with_issues(self):
        """Create a docx Document with multiple contrast issues."""
        import docx
        from docx.shared import RGBColor as RGB

        doc = docx.Document()
        # First paragraph: gray text
        p1 = doc.add_paragraph()
        r1 = p1.add_run("Light gray text")
        r1.font.color.rgb = RGB(0xAA, 0xAA, 0xAA)
        # Second paragraph: another gray run
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Another gray run")
        r2.font.color.rgb = RGB(0xBB, 0xBB, 0xBB)

        # Matching ParagraphInfo models
        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="Light gray text",
                runs=[RunInfo(text="Light gray text", color="#AAAAAA", font_size_pt=12.0)],
            ),
            ParagraphInfo(
                id="p_1",
                text="Another gray run",
                runs=[RunInfo(text="Another gray run", color="#BBBBBB", font_size_pt=12.0)],
            ),
        ]
        return doc, paragraphs

    def test_fixes_all_issues(self):
        doc, paragraphs = self._make_doc_with_issues()
        result = fix_all_document_contrast(doc, paragraphs)
        assert result.fixes_applied == 2
        assert result.fixes_failed == 0
        assert len(result.changes) == 2

        # Verify both colors now pass contrast
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    fg_hex = f"#{run.font.color.rgb}"
                    check = check_contrast(fg_hex, "#FFFFFF")
                    assert check.passes, f"Color {fg_hex} still fails contrast"

    def test_no_issues_returns_success(self):
        import docx

        doc = docx.Document()
        doc.add_paragraph("Black text")  # default black = passes

        paragraphs = [
            ParagraphInfo(
                id="p_0",
                text="Black text",
                runs=[RunInfo(text="Black text", font_size_pt=12.0)],
            ),
        ]
        result = fix_all_document_contrast(doc, paragraphs)
        assert result.success
        assert result.fixes_applied == 0
        assert "No contrast issues" in result.changes[0]
