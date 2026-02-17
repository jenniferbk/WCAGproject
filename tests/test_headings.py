"""Tests for heading detection and manipulation."""

from pathlib import Path

import docx
import pytest

from src.models.document import FakeHeadingSignals, ParagraphInfo
from src.tools.docx_parser import parse_docx
from src.tools.headings import (
    get_fake_heading_candidates,
    set_heading_level,
    suggest_heading_level,
    validate_heading_hierarchy,
)


class TestValidateHeadingHierarchy:
    def test_valid_hierarchy(self, simple_docx: Path):
        result = parse_docx(simple_docx)
        issues = validate_heading_hierarchy(result.document.paragraphs)
        assert len(issues) == 0

    def test_skipped_level(self, skipped_headings_docx: Path):
        result = parse_docx(skipped_headings_docx)
        issues = validate_heading_hierarchy(result.document.paragraphs)
        skipped = [i for i in issues if i.issue_type == "skipped_level"]
        assert len(skipped) >= 1
        assert "skips from" in skipped[0].detail

    def test_no_h1(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="Section", heading_level=2),
            ParagraphInfo(id="p_1", text="Body"),
        ]
        issues = validate_heading_hierarchy(paragraphs)
        no_h1 = [i for i in issues if i.issue_type == "no_h1"]
        assert len(no_h1) == 1

    def test_multiple_h1(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="First Title", heading_level=1),
            ParagraphInfo(id="p_1", text="Body"),
            ParagraphInfo(id="p_2", text="Second Title", heading_level=1),
        ]
        issues = validate_heading_hierarchy(paragraphs)
        multi = [i for i in issues if i.issue_type == "multiple_h1"]
        assert len(multi) == 1


class TestGetFakeHeadingCandidates:
    def test_finds_candidates(self, fake_headings_docx: Path):
        result = parse_docx(fake_headings_docx)
        candidates = get_fake_heading_candidates(result.document.paragraphs)
        assert len(candidates) >= 1
        texts = [c[0].text for c in candidates]
        assert any("Fake Section Title" in t for t in texts)

    def test_min_score_filter(self, fake_headings_docx: Path):
        result = parse_docx(fake_headings_docx)
        low = get_fake_heading_candidates(result.document.paragraphs, min_score=0.3)
        high = get_fake_heading_candidates(result.document.paragraphs, min_score=0.8)
        assert len(low) >= len(high)


class TestSetHeadingLevel:
    def test_convert_to_heading(self, fake_headings_docx: Path):
        doc = docx.Document(str(fake_headings_docx))

        # Find "Fake Section Title" paragraph index
        for i, para in enumerate(doc.paragraphs):
            if "Fake Section Title" in para.text:
                break

        result = set_heading_level(doc, i, 2)
        assert result.success
        assert doc.paragraphs[i].style.name == "Heading 2"

    def test_clears_bold(self, fake_headings_docx: Path):
        doc = docx.Document(str(fake_headings_docx))
        for i, para in enumerate(doc.paragraphs):
            if "Fake Section Title" in para.text:
                break

        set_heading_level(doc, i, 2)
        # Bold should be cleared (None = inherit from style)
        for run in doc.paragraphs[i].runs:
            assert run.bold is None

    def test_invalid_level(self, simple_docx: Path):
        doc = docx.Document(str(simple_docx))
        result = set_heading_level(doc, 0, 0)
        assert not result.success

    def test_invalid_index(self, simple_docx: Path):
        doc = docx.Document(str(simple_docx))
        result = set_heading_level(doc, 999, 1)
        assert not result.success


class TestSuggestHeadingLevel:
    def test_first_heading_suggests_1(self):
        para = ParagraphInfo(
            id="p_0", text="Title",
            fake_heading_signals=FakeHeadingSignals(score=0.8),
        )
        level = suggest_heading_level(para, [])
        assert level == 1

    def test_after_h1_suggests_2(self):
        para = ParagraphInfo(
            id="p_2", text="Section",
            fake_heading_signals=FakeHeadingSignals(
                font_size_pt=14.0, score=0.7
            ),
        )
        level = suggest_heading_level(para, [("p_0", 1)])
        assert level == 2
