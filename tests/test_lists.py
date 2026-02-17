"""Tests for fake list detection and conversion."""

from pathlib import Path

import docx
import pytest

from src.models.document import ParagraphInfo
from src.tools.docx_parser import parse_docx
from src.tools.lists import ListType, convert_to_list, detect_fake_lists


class TestDetectFakeLists:
    def test_detect_numbered_list(self, fake_lists_docx: Path):
        result = parse_docx(fake_lists_docx)
        candidates = detect_fake_lists(result.document.paragraphs)

        numbered = [c for c in candidates if c.list_type == ListType.NUMBERED]
        assert len(numbered) >= 1
        assert len(numbered[0].paragraph_ids) == 3

    def test_detect_bulleted_list(self, fake_lists_docx: Path):
        result = parse_docx(fake_lists_docx)
        candidates = detect_fake_lists(result.document.paragraphs)

        bulleted = [c for c in candidates if c.list_type == ListType.BULLETED]
        assert len(bulleted) >= 1
        assert len(bulleted[0].paragraph_ids) == 3

    def test_needs_at_least_two_items(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="1. Only one numbered item"),
            ParagraphInfo(id="p_1", text="Regular text."),
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 0

    def test_skips_real_list_items(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="1. First", is_list_item=True),
            ParagraphInfo(id="p_1", text="2. Second", is_list_item=True),
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 0

    def test_skips_headings(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="1. First", heading_level=1),
            ParagraphInfo(id="p_1", text="2. Second", heading_level=2),
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 0

    def test_mixed_types_separate(self):
        paragraphs = [
            ParagraphInfo(id="p_0", text="1. First"),
            ParagraphInfo(id="p_1", text="2. Second"),
            ParagraphInfo(id="p_2", text="- Bullet one"),
            ParagraphInfo(id="p_3", text="- Bullet two"),
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 2

    def test_confidence_increases_with_length(self):
        paragraphs = [
            ParagraphInfo(id=f"p_{i}", text=f"{i+1}. Item {i+1}")
            for i in range(5)
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 1
        assert candidates[0].confidence > 0.5

    def test_bullet_variations(self):
        """Various bullet characters should be detected."""
        paragraphs = [
            ParagraphInfo(id="p_0", text="\u2022 Unicode bullet"),
            ParagraphInfo(id="p_1", text="\u2022 Another bullet"),
        ]
        candidates = detect_fake_lists(paragraphs)
        assert len(candidates) == 1
        assert candidates[0].list_type == ListType.BULLETED


class TestConvertToList:
    def test_convert_to_numbered(self, fake_lists_docx: Path):
        doc = docx.Document(str(fake_lists_docx))

        # Find the numbered list paragraphs
        indices = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.startswith(("1.", "2.", "3.")):
                indices.append(i)

        result = convert_to_list(doc, indices, ListType.NUMBERED)
        assert result.success
        assert len(result.changes) == 3

        # Verify style was applied
        for idx in indices:
            assert "List Number" in doc.paragraphs[idx].style.name

    def test_convert_to_bulleted(self, fake_lists_docx: Path):
        doc = docx.Document(str(fake_lists_docx))

        indices = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.startswith("- "):
                indices.append(i)

        result = convert_to_list(doc, indices, ListType.BULLETED)
        assert result.success

        for idx in indices:
            assert "List Bullet" in doc.paragraphs[idx].style.name

    def test_strips_prefix(self, fake_lists_docx: Path):
        doc = docx.Document(str(fake_lists_docx))

        indices = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.startswith("1."):
                indices.append(i)
                break

        convert_to_list(doc, indices, ListType.NUMBERED, strip_prefix=True)
        assert not doc.paragraphs[indices[0]].text.startswith("1.")

    def test_invalid_index(self, fake_lists_docx: Path):
        doc = docx.Document(str(fake_lists_docx))
        result = convert_to_list(doc, [999], ListType.BULLETED)
        assert not result.success
