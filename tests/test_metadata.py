"""Tests for metadata tools."""

from pathlib import Path

import docx
import pytest

from src.tools.metadata import fix_metadata, set_language, set_title


class TestSetTitle:
    def test_set_title(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = set_title(doc, "Introduction to Biology")
        assert result.success
        assert doc.core_properties.title == "Introduction to Biology"

    def test_overwrite_title(self, simple_docx: Path):
        doc = docx.Document(str(simple_docx))
        result = set_title(doc, "New Title")
        assert result.success
        assert doc.core_properties.title == "New Title"
        assert "Test Document" in result.changes[0]

    def test_empty_title_rejected(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = set_title(doc, "")
        assert not result.success

    def test_whitespace_only_rejected(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = set_title(doc, "   ")
        assert not result.success

    def test_strips_whitespace(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        set_title(doc, "  Padded Title  ")
        assert doc.core_properties.title == "Padded Title"


class TestSetLanguage:
    def test_set_language(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = set_language(doc, "en-US")
        assert result.success
        assert doc.core_properties.language == "en-US"

    def test_empty_language_rejected(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = set_language(doc, "")
        assert not result.success


class TestFixMetadata:
    def test_fix_missing_both(self, no_metadata_docx: Path):
        doc = docx.Document(str(no_metadata_docx))
        result = fix_metadata(doc, title="Course Syllabus", language="en")
        assert result.success
        assert doc.core_properties.title == "Course Syllabus"
        assert doc.core_properties.language == "en"
        assert len(result.changes) == 2

    def test_fix_only_language_when_title_exists(self, simple_docx: Path):
        doc = docx.Document(str(simple_docx))
        # simple_docx has title but let's clear language
        doc.core_properties.language = ""
        result = fix_metadata(doc, title="Fallback", language="en")
        assert result.success
        # Title should NOT be overwritten (it already exists)
        assert doc.core_properties.title == "Test Document"

    def test_no_changes_needed(self, simple_docx: Path):
        doc = docx.Document(str(simple_docx))
        result = fix_metadata(doc, title="Fallback", language="en")
        assert result.success
        assert "No metadata changes needed" in result.changes[0]

    def test_persists_after_save(self, no_metadata_docx: Path, tmp_path: Path):
        doc = docx.Document(str(no_metadata_docx))
        fix_metadata(doc, title="Saved Title", language="fr")

        output = tmp_path / "saved.docx"
        doc.save(str(output))

        doc2 = docx.Document(str(output))
        assert doc2.core_properties.title == "Saved Title"
        assert doc2.core_properties.language == "fr"
