"""Tests for table header marking and analysis."""

from pathlib import Path

import docx
import pytest
from docx.oxml.ns import qn

from src.models.document import CellInfo, TableInfo
from src.tools.docx_parser import parse_docx
from src.tools.tables import (
    analyze_all_tables,
    analyze_table,
    mark_header_rows,
    remove_header_rows,
)


class TestAnalyzeTable:
    def test_table_with_header(self, table_docx: Path):
        result = parse_docx(table_docx)
        analysis = analyze_table(result.document.tables[0])
        assert analysis.has_header_row
        assert analysis.row_count == 3
        assert analysis.col_count == 3
        assert len(analysis.issues) == 0

    def test_table_without_header(self):
        table = TableInfo(
            id="tbl_0",
            rows=[[CellInfo(text="A"), CellInfo(text="B")]],
            header_row_count=0,
            row_count=1,
            col_count=2,
        )
        analysis = analyze_table(table)
        assert not analysis.has_header_row
        assert any("no marked header" in i for i in analysis.issues)

    def test_merged_cell_flagged(self):
        table = TableInfo(
            id="tbl_0",
            rows=[
                [CellInfo(text="A", grid_span=2), CellInfo(text="B")],
                [CellInfo(text="C"), CellInfo(text="D"), CellInfo(text="E")],
            ],
            header_row_count=1,
            row_count=2,
            col_count=3,
        )
        analysis = analyze_table(table)
        assert analysis.has_merged_cells
        assert analysis.is_complex
        assert any("merged" in i.lower() for i in analysis.issues)


class TestMarkHeaderRows:
    def test_mark_header(self, table_no_header_docx: Path):
        doc = docx.Document(str(table_no_header_docx))

        result = mark_header_rows(doc, 0, header_count=1)
        assert result.success

        # Verify XML
        row = doc.tables[0].rows[0]
        tr_pr = row._tr.find(qn("w:trPr"))
        assert tr_pr is not None
        assert tr_pr.find(qn("w:tblHeader")) is not None

    def test_already_marked(self, table_docx: Path):
        doc = docx.Document(str(table_docx))
        result = mark_header_rows(doc, 0, header_count=1)
        assert result.success
        assert any("already" in c for c in result.changes)

    def test_invalid_table_index(self, table_docx: Path):
        doc = docx.Document(str(table_docx))
        result = mark_header_rows(doc, 99)
        assert not result.success
        assert "out of range" in result.error

    def test_header_count_exceeds_rows(self, table_docx: Path):
        doc = docx.Document(str(table_docx))
        result = mark_header_rows(doc, 0, header_count=100)
        assert not result.success

    def test_persists_after_save(self, table_no_header_docx: Path, tmp_path: Path):
        doc = docx.Document(str(table_no_header_docx))
        mark_header_rows(doc, 0, header_count=1)

        output = tmp_path / "saved.docx"
        doc.save(str(output))

        doc2 = docx.Document(str(output))
        row = doc2.tables[0].rows[0]
        tr_pr = row._tr.find(qn("w:trPr"))
        assert tr_pr is not None
        assert tr_pr.find(qn("w:tblHeader")) is not None


class TestRemoveHeaderRows:
    def test_remove_header(self, table_docx: Path):
        doc = docx.Document(str(table_docx))
        result = remove_header_rows(doc, 0)
        assert result.success

        row = doc.tables[0].rows[0]
        tr_pr = row._tr.find(qn("w:trPr"))
        if tr_pr is not None:
            assert tr_pr.find(qn("w:tblHeader")) is None

    def test_remove_no_headers(self, table_no_header_docx: Path):
        doc = docx.Document(str(table_no_header_docx))
        result = remove_header_rows(doc, 0)
        assert result.success
        assert any("no header" in c for c in result.changes)


class TestAnalyzeAllTables:
    def test_multiple_tables(self):
        tables = [
            TableInfo(id="tbl_0", rows=[], header_row_count=1, row_count=2, col_count=2),
            TableInfo(id="tbl_1", rows=[], header_row_count=0, row_count=3, col_count=3),
        ]
        analyses = analyze_all_tables(tables)
        assert len(analyses) == 2
        assert analyses[0].has_header_row
        assert not analyses[1].has_header_row
